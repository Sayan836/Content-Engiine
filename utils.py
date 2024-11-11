import fitz  # PyMuPDF
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import docx 

def extract_pdf_to_documents(pdf_file):
    """
    Extracts text content from an uploaded PDF file object and returns a list of Document objects.
    
    Args:
        pdf_file: A file-like object representing the uploaded PDF (e.g., from Streamlit's file_uploader).
    
    Returns:
        List of Document objects with text content from each page.
    """
    documents = []
    
    # Open the PDF file using a file-like object
    with fitz.open(stream=pdf_file.read(), filetype="pdf") as pdf:
        for page_num in range(pdf.page_count):
            page = pdf.load_page(page_num)
            text = page.get_text()
            # Create a Document object for each page with metadata
            documents.append(Document(page_content=text, metadata={"page": page_num + 1}))

    return documents

# split extracted_text
def split_text(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=2500, chunk_overlap=200)
    docs= text_splitter.split_documents(text)
    print("Total number of documents' chunks: ",len(docs))
    return docs

def extract_docx_to_documents(docx_path):
    """Extract text from a DOCX file and return as a list of Document objects."""
    documents = []
    doc = docx.Document(docx_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    text = "\n".join(full_text)
    documents.append(Document(page_content=text, metadata={"file_type": "docx"}))
    return documents




